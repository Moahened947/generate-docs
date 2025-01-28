import os
import json
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import traceback

# Configure Google Gemini API
GOOGLE_API_KEY = "AIzaSyAOdKEsip0P1DkWvC6tMO4hl3jG0rwZc7Y"
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize Gemini model
model = genai.GenerativeModel('gemini-pro')

def clean_text(text):
    """Clean text by removing markdown symbols and normalizing line breaks."""
    # Remove markdown bold symbols
    text = text.replace('**', '')
    # Normalize line breaks and indentation
    lines = [line.strip() for line in text.split('\n')]
    return ' '.join(lines)

def set_rtl_font(run):
    """Set RTL font properties for Arabic text."""
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')

def create_document(content_json, filename, lang='ar'):
    """Create a Word document with the generated content."""
    doc = Document()
    
    # Set RTL for Arabic
    if lang == 'ar':
        # Enable RTL for the document
        section = doc.sections[0]._sectPr
        section.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
        
        # Set RTL for default style
        style = doc.styles['Normal']
        style._element.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
        style._element.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection", "rtl")
        style.font.size = Pt(12)  # Set default font size to 12pt
        
        # Set RTL for heading styles
        for i in range(1, 10):
            try:
                heading_style = doc.styles[f'Heading {i}']
                heading_style._element.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
                heading_style._element.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection", "rtl")
                heading_style.font.size = Pt(14)  # Set heading font size to 14pt
            except KeyError:
                continue
    
    # Add title
    title = doc.add_heading(content_json['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set title font size to 16pt
    for run in title.runs:
        run.font.size = Pt(16)
    if lang == 'ar':
        title._p.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
        title._p.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection", "rtl")
        # Set RTL for runs in title
        for run in title.runs:
            run._element.get_or_add_rPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rtl", "1")
    
    # Add sections
    for section in content_json['sections']:
        # Add section heading
        heading = doc.add_heading(section['heading'], 1)
        # Set heading font size to 14pt
        for run in heading.runs:
            run.font.size = Pt(14)
        if lang == 'ar':
            heading._p.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
            heading._p.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection", "rtl")
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Set RTL for runs in heading
            for run in heading.runs:
                run._element.get_or_add_rPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rtl", "1")
        
        # Process content and handle bullet points
        content = section['content']
        paragraphs = content.split('\n')
        
        for para_text in paragraphs:
            # Skip empty paragraphs
            if not para_text.strip():
                continue
                
            # Check if this is a bullet point
            if para_text.strip().startswith('*') or para_text.strip().startswith('-'):
                # Remove the bullet point marker and any leading/trailing whitespace
                text = para_text.strip().lstrip('*').lstrip('-').strip()
                # Add as a bullet point
                p = doc.add_paragraph(text, style='List Bullet')
            else:
                # Add as normal paragraph
                p = doc.add_paragraph(para_text)
            
            # Set content font size to 12pt
            for run in p.runs:
                run.font.size = Pt(12)
            
            if lang == 'ar':
                p._p.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bidi", "1")
                p._p.get_or_add_pPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection", "rtl")
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # Set RTL for runs in paragraph
                for run in p.runs:
                    run._element.get_or_add_rPr().set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rtl", "1")
    
    doc.save(filename)
    return f"Document saved successfully as {filename}"

def generate_research_content(topic, lang='ar', doc_type='academic', instructions=''):
    """Generate research content using Google Gemini."""
    
    # تحديد نوع المستند بالعربية
    doc_types_ar = {
        'academic': 'بحث علمي',
        'technical_report': 'تقرير فني',
        'administrative_report': 'تقرير إداري',
        'scientific_report': 'تقرير علمي',
        'financial_report': 'تقرير مالي',
        'project_report': 'تقرير مشروع'
    }
    
    # تحديد نوع المستند بالإنجليزية
    doc_types_en = {
        'academic': 'Academic Research',
        'technical_report': 'Technical Report',
        'administrative_report': 'Administrative Report',
        'scientific_report': 'Scientific Report',
        'financial_report': 'Financial Report',
        'project_report': 'Project Report'
    }
    
    # تحديد نوع المستند حسب اللغة
    doc_type_name = doc_types_ar[doc_type] if lang == 'ar' else doc_types_en[doc_type]
    
    if lang == 'ar':
        default_prompt = f"""
قم بإعداد {doc_type_name} باللغة العربية حول: {topic}

المطلوب تنسيق JSON كالتالي:
{{
    "title": "عنوان المستند",
    "sections": [
        {{
            "heading": "عنوان القسم",
            "content": "محتوى القسم"
        }}
    ]
}}

متطلبات المحتوى:
- قم بتحديد الأقسام المناسبة لنوع المستند ({doc_type_name})
- كل قسم يجب أن يحتوي على محتوى لا يقل عن 500 كلمة
- استخدم لغة مهنية واضحة
- قم بدعم الأفكار بالأدلة والمراجع
- اجعل المحتوى شاملاً ومفصلاً
- تأكد من الترابط المنطقي بين الأقسام
"""
    else:
        default_prompt = f"""
Create a {doc_type_name} in English about: {topic}

Required JSON format:
{{
    "title": "Document Title",
    "sections": [
        {{
            "heading": "Section Heading",
            "content": "Section Content"
        }}
    ]
}}

Content requirements:
- Determine appropriate sections for the document type ({doc_type_name})
- Each section should contain at least 500 words
- Use clear professional language
- Support ideas with evidence and references
- Make the content comprehensive and detailed
- Ensure logical connection between sections
"""

    if instructions:
        if lang == 'ar':
            prompt = f"""
قم بإعداد {doc_type_name} باللغة العربية حول: {topic}

متطلبات خاصة من المستخدم:
{instructions}

المطلوب تنسيق JSON كالتالي:
{{
    "title": "عنوان المستند",
    "sections": [
        {{
            "heading": "عنوان القسم",
            "content": "محتوى القسم"
        }}
    ]
}}

متطلبات المحتوى:
- قم بتحديد الأقسام المناسبة لنوع المستند ({doc_type_name})
- كل قسم يجب أن يحتوي على محتوى لا يقل عن 500 كلمة
- استخدم لغة مهنية واضحة
- قم بدعم الأفكار بالأدلة والمراجع
- اجعل المحتوى شاملاً ومفصلاً
- تأكد من الترابط المنطقي بين الأقسام
"""
        else:
            prompt = f"""
Create a {doc_type_name} in English about: {topic}

Special requirements from user:
{instructions}

Required JSON format:
{{
    "title": "Document Title",
    "sections": [
        {{
            "heading": "Section Heading",
            "content": "Section Content"
        }}
    ]
}}

Content requirements:
- Determine appropriate sections for the document type ({doc_type_name})
- Each section should contain at least 500 words
- Use clear professional language
- Support ideas with evidence and references
- Make the content comprehensive and detailed
- Ensure logical connection between sections
"""
    else:
        prompt = default_prompt

    try:
        # Add error handling and logging
        print(f"Generating content for topic: {topic}, language: {lang}, type: {doc_type}")
        
        # Generate content
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Log the response for debugging
        print("Raw response received:", response_text[:100] + "...")
        
        # Remove any markdown code block indicators if present
        response_text = response_text.replace('```json', '').replace('```', '').strip()
        
        # First attempt to parse JSON
        try:
            content = json.loads(response_text)
            print("Successfully parsed JSON response")
            return content
        except json.JSONDecodeError as e:
            print(f"Initial JSON parsing failed: {e}")
            
            # If direct parsing fails, try to clean the response further
            print("Attempting to clean the response...")
            
            # Replace problematic characters
            cleaned_text = response_text.replace('\n', ' ').replace('\r', ' ')
            cleaned_text = ' '.join(cleaned_text.split())
            
            # Try parsing again
            try:
                content = json.loads(cleaned_text)
                print("Successfully parsed JSON after cleaning")
                return content
            except json.JSONDecodeError as e:
                print(f"Failed to parse JSON after cleaning: {e}")
                print("Cleaned response:", cleaned_text)
                return None
            
    except Exception as e:
        print(f"Error generating content: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        if hasattr(e, '__traceback__'):
            print("Traceback:", traceback.format_exc())
        return None

def main():
    topic = input("Enter the research topic: ")
    lang = input("Enter the language (ar/en): ")
    doc_type = input("Enter the document type (academic/report): ")
    instructions = input("Enter any special instructions (optional): ")
    filename = f"research_{topic.replace(' ', '_').lower()}_{lang}_{doc_type}.docx"
    
    print(f"\nGenerating research content about '{topic}' in {lang}...")
    content = generate_research_content(topic, lang, doc_type, instructions)
    
    if content:
        print("\nCreating Word document...")
        result = create_document(content, filename, lang)
        print(result)
    else:
        print("Failed to generate content. Please try again.")

if __name__ == "__main__":
    main()
